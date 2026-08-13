from pathlib import Path
import re

from faker import Faker
from docx import Document


fake = Faker("en_IN")

# Make generated values reproducible.
Faker.seed(42)


def generate_fake_value(pii_type: str) -> str:
    """
    Generate a realistic fake replacement for a PII category.
    """

    if pii_type == "PERSON":
        return fake.name()

    if pii_type == "EMAIL":
        return fake.email()

    if pii_type == "PHONE":
        return fake.phone_number()

    if pii_type == "COMPANY":
        return fake.company()

    if pii_type == "SSN":
        return fake.ssn()

    if pii_type == "CREDIT_CARD":
        return fake.credit_card_number()

    if pii_type == "DOB":
        return fake.date_of_birth().strftime("%d/%m/%Y")

    if pii_type == "IP_ADDRESS":
        return fake.ipv4()

    if pii_type == "ADDRESS":
        return fake.address().replace("\n", ", ")

    return "[REDACTED]"


def build_replacement_map(
    detections: dict[str, list[str]],
) -> dict[str, str]:
    """
    Create one fake replacement for every detected PII value.

    The same original value always receives the same replacement.
    """

    replacements = {}

    for pii_type, values in detections.items():
        for value in values:
            if not value:
                continue

            if value not in replacements:
                replacements[value] = generate_fake_value(pii_type)

    return replacements


def _replace_in_paragraph(paragraph, replacements: dict[str, str]) -> None:
    """
    Replace detected PII inside a paragraph.

    Handles both:
      - PII contained inside one run
      - PII split across multiple Word runs

    Existing formatting is retained as much as possible.
    """

    if not paragraph.runs:
        return

    original_text = "".join(run.text or "" for run in paragraph.runs)

    if not original_text:
        return

    matches = []

    ordered_replacements = sorted(
        replacements.items(),
        key=lambda item: len(item[0]),
        reverse=True,
    )

    for original, replacement in ordered_replacements:
        if not original:
            continue

        start = 0

        while True:
            index = original_text.casefold().find(
                original.casefold(),
                start,
            )

            if index == -1:
                break

            matches.append(
                (
                    index,
                    index + len(original),
                    replacement,
                )
            )

            start = index + len(original)

    if not matches:
        return

    # Remove overlapping matches.
    matches.sort(key=lambda item: (item[0], -(item[1] - item[0])))

    selected = []

    for match in matches:
        start, end, replacement = match

        if selected and start < selected[-1][1]:
            continue

        selected.append(match)

    for start, end, replacement in reversed(selected):
        _replace_character_range(
            paragraph,
            start,
            end,
            replacement,
        )


def _replace_character_range(
    paragraph,
    start: int,
    end: int,
    replacement: str,
) -> None:
    """
    Replace a character range spanning one or more runs.
    """

    runs = paragraph.runs

    positions = []
    current = 0

    for index, run in enumerate(runs):
        text = run.text or ""
        run_start = current
        run_end = current + len(text)

        positions.append(
            (
                index,
                run_start,
                run_end,
            )
        )

        current = run_end

    start_run = None
    end_run = None

    for index, run_start, run_end in positions:
        if run_start <= start < run_end:
            start_run = index

        if run_start < end <= run_end:
            end_run = index
            break

    if start_run is None:
        return

    if end == start:
        return

    if end_run is None:
        end_run = len(runs) - 1

    start_run_text = runs[start_run].text or ""

    start_offset = start - positions[start_run][1]

    if start_run == end_run:
        end_offset = end - positions[end_run][1]

        runs[start_run].text = (
            start_run_text[:start_offset]
            + replacement
            + start_run_text[end_offset:]
        )

        return

    end_run_text = runs[end_run].text or ""
    end_offset = end - positions[end_run][1]

    prefix = start_run_text[:start_offset]
    suffix = end_run_text[end_offset:]

    runs[start_run].text = prefix + replacement

    for index in range(start_run + 1, end_run):
        runs[index].text = ""

    runs[end_run].text = suffix


def _replace_in_table(
    table,
    replacements: dict[str, str],
) -> None:
    """
    Recursively process a Word table.
    """

    for row in table.rows:
        for cell in row.cells:
            _replace_in_container(
                cell,
                replacements,
            )


def _replace_in_container(
    container,
    replacements: dict[str, str],
) -> None:
    """
    Process paragraphs and nested tables.
    """

    for paragraph in container.paragraphs:
        _replace_in_paragraph(
            paragraph,
            replacements,
        )

    for table in container.tables:
        _replace_in_table(
            table,
            replacements,
        )


def redact_docx(
    input_file: str | Path,
    output_file: str | Path,
    detections: dict[str, list[str]],
) -> dict[str, str]:
    """
    Redact detected PII from a DOCX document.

    Returns the replacement mapping used.
    """

    input_file = Path(input_file)
    output_file = Path(output_file)

    document = Document(input_file)

    replacements = build_replacement_map(detections)

    # Main document body.
    _replace_in_container(
        document,
        replacements,
    )

    # Headers and footers.
    for section in document.sections:
        _replace_in_container(
            section.header,
            replacements,
        )

        _replace_in_container(
            section.footer,
            replacements,
        )

    output_file.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    document.save(output_file)

    return replacements