from pathlib import Path

from docx import Document, text


def extract_table_text(table) -> list[str]:
    """pakdta non-empty text from every cell in a Word table."""
    rows = []

    for row in table.rows:
        cells = []

        for cell in row.cells:
            text = cell.text.strip()

            if text:
                cells.append(text)

        if cells:
            rows.append(" | ".join(cells))

    return rows


def read_docx(file_path: Path) -> str:
    """Read docx file and vapis krta text ko string mein return karta hai."""
    document = Document(file_path)

    content = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()

        if text:
            content.append(text)

    
    for table in document.tables:
        table_rows = extract_table_text(table)
        content.extend(table_rows)

    return "\n".join(content)


def main() -> None:
    input_file = Path("input/Red Herring Prospectus.docx")

    if not input_file.exists():
        raise FileNotFoundError(
            f"Input document not found: {input_file}"
        )

    document_text = read_docx(input_file)

    print("Document loaded successfully.")
    print(f"Characters extracted: {len(document_text)}")
    print(f"Lines extracted: {len(document_text.splitlines())}")

    print("\nFirst 1000 characters:\n")
    print(document_text[:1000])

    print("\nTable extraction check:")

    keywords = [
        "Kushal Subbayya Hegde",
        "Rajesh Kushal Hegde",
        "Rohit Kushal Hegde",
    ]

    for keyword in keywords:
        if keyword in document_text:
            print(f"[FOUND] {keyword}")
        else:
            print(f"[NOT FOUND] {keyword}")



if __name__ == "__main__":
    main()

